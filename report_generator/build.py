"""Assemble the Word research note (python-docx) from the structured inputs, the DCF
truth, the narrative, and the charts. Layout follows a standard sell-side note:
front page (masthead · rating & data rail · thesis · financial summary), key highlights
(interim results · valuation summary · chart), narrative, financial highlights, disclaimer.

Run: uv run python -m report_generator.build
"""
from __future__ import annotations

import json

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from pipeline import config
from report_generator import charts, narrative

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x44, 0x72, 0xC4)
GREEN = RGBColor(0x1E, 0x7A, 0x34)
RED = RGBColor(0xB0, 0x20, 0x20)
GREY = RGBColor(0x59, 0x59, 0x59)


# ── low-level helpers ─────────────────────────────────────────────────────
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def run(p, text, *, size=10, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def para(doc, text="", *, size=10, bold=False, color=None, before=2, after=2, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        run(p, text, size=size, bold=bold, color=color)
    return p


def heading(doc, text):
    p = para(doc, text, size=12, bold=True, color=NAVY, before=10, after=4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1F3864")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def money(x, dp=0):
    return f"({abs(x):,.{dp}f})" if x < 0 else f"{x:,.{dp}f}"


# ── context ───────────────────────────────────────────────────────────────
def load_context() -> dict:
    import yaml
    d = config.output_dir_for()
    ctx = {
        "meta": config.FILING_META,
        "financials": json.load(open(d / "financials.json")),
        "model_values": json.load(open(d / "model_values.json")),
        "market": yaml.safe_load((config.CONFIG_DIR / "market_data.yaml").read_text()),
    }
    for opt in ("commentary", "news"):
        f = d / f"{opt}.json"
        ctx[opt] = json.load(open(f)) if f.exists() else {}
    v = ctx["model_values"]["valuation"]
    ctx["mktcap"] = v["current_price"] * v["common_shares"] / 1e9  # ₱bn
    return ctx


# ── sections ──────────────────────────────────────────────────────────────
def _front(doc, ctx, nar):
    m, mv, mk = ctx["meta"], ctx["model_values"], ctx["market"]
    v = mv["valuation"]

    # masthead
    t = doc.add_table(rows=1, cols=2); no_borders(t)
    run(t.cell(0, 0).paragraphs[0], "EQUITY RESEARCH", size=12, bold=True, color=NAVY)
    pr = t.cell(0, 1).paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(pr, f"Philippines · Utilities\n{v['as_of_date'] if 'as_of_date' in v else mv['as_of_date']}", size=9, color=GREY)

    para(doc, f"{m['company']}", size=18, bold=True, color=NAVY, before=4, after=0)
    para(doc, f"PSE: {m['ticker']}   ·   Bloomberg: {m['bloomberg']}   ·   Regulated water utility",
         size=9.5, color=GREY, after=6)

    # rail: left = rating + thesis, right = data box
    rail = doc.add_table(rows=1, cols=2); no_borders(t)
    rail.columns[0].width = Inches(4.6); rail.columns[1].width = Inches(2.2)
    left, right = rail.cell(0, 0), rail.cell(0, 1)

    rec_color = GREEN if v["recommendation"] == "BUY" else (GREY if v["recommendation"] == "HOLD" else RED)
    p = left.paragraphs[0]
    run(p, f"{v['recommendation']}", size=20, bold=True, color=rec_color)
    run(p, f"   Target ₱{v['fair_value_per_share']:.2f}", size=13, bold=True)
    run(p, f"   ({v['upside_pct']:+.0%})", size=12, bold=True, color=rec_color)
    rp = left.add_paragraph(); rp.paragraph_format.space_after = Pt(6)
    run(rp, "Investment thesis", size=10.5, bold=True, color=NAVY)
    for b in nar["thesis_bullets"]:
        bp = left.add_paragraph(style="List Bullet"); bp.paragraph_format.space_after = Pt(3)
        lead, _, rest = b.partition("—")
        run(bp, lead.strip() + (" — " if rest else ""), size=9.5, bold=True)
        run(bp, rest.strip(), size=9.5)

    shade(right, "EEF1F8")
    dp = right.paragraphs[0]; run(dp, "Company data", size=10, bold=True, color=NAVY)
    data = [
        ("Price (₱)", f"{v['current_price']:.2f}"),
        ("Fair value (₱)", f"{v['fair_value_per_share']:.2f}"),
        ("Upside", f"{v['upside_pct']:+.0%}"),
        ("Market cap (₱bn)", f"{ctx['mktcap']:.1f}"),
        ("52-wk range (₱)", f"{mk.get('week52_low','—')}–{mk.get('week52_high','—')}"),
        ("Shares (m)", f"{v['common_shares']/1e6:,.0f}"),
        ("Beta", f"{mk.get('beta','—')}"),
        ("WACC", f"{mv['wacc']['wacc']:.1%}"),
    ]
    for k, val in data:
        pp = right.add_paragraph(); pp.paragraph_format.space_after = Pt(1)
        run(pp, f"{k}: ", size=9, color=GREY); run(pp, val, size=9, bold=True)

    # financial summary mini-table
    para(doc, "", after=2)
    _fin_summary(doc, ctx)


def _fin_summary(doc, ctx):
    p = ctx["model_values"]["projections"]
    yrs = [2024, 2025, 2026, 2027]
    rows = [("₱m unless stated", [f"{y}{'E' if y<=2027 else 'F'}" for y in yrs])]
    rows.append(("Revenue", [money(p["revenue"][str(y)] / 1000) for y in yrs]))
    rows.append(("Net income (parent)", [money(p["net_income"][str(y)] / 1000) for y in yrs]))
    rows.append(("EPS (₱)", [f"{p['eps'][str(y)]:.2f}" for y in yrs]))
    eps = [p["eps"][str(y)] for y in yrs]
    rows.append(("EPS growth", ["—"] + [f"{eps[i]/eps[i-1]-1:+.0%}" for i in range(1, len(eps))]))
    _grid(doc, rows, first_w=2.3)


def _grid(doc, rows, first_w=2.6, header=True):
    nval = len(rows[0][1])
    ncol = 1 + nval
    usable = 7.2                       # page width minus margins
    val_w = (usable - first_w) / nval
    widths = [Inches(first_w)] + [Inches(val_w)] * nval
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    t.autofit = False
    t.allow_autofit = False
    for i, (label, vals) in enumerate(rows):
        cells = t.rows[i].cells
        for ci, c in enumerate(cells):
            c.width = widths[ci]
        lp = cells[0].paragraphs[0]
        run(lp, label, size=9, bold=(i == 0))
        if i == 0 and header:
            shade(cells[0], "1F3864"); lp.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for j, val in enumerate(vals):
            cp = cells[j + 1].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run(cp, str(val), size=9, bold=(i == 0),
                color=(RGBColor(0xFF, 0xFF, 0xFF) if i == 0 and header else None))
            if i == 0 and header:
                shade(cells[j + 1], "1F3864")
    return t


def _key_highlights(doc, ctx):
    doc.add_page_break()
    heading(doc, "Key highlights")
    mv = ctx["model_values"]; v = mv["valuation"]

    para(doc, "Interim results (₱m)", size=10.5, bold=True, color=NAVY, before=6, after=3)
    F = ctx["financials"]
    inc = {li["account"]: li for li in F["income"]["line_items"] if li["account"]}
    want = [("revenue_total", "Revenue"), ("gross_profit", "Gross profit"),
            ("income_before_other", "Operating income"), ("income_before_tax", "Pre-tax income"),
            ("net_income_total", "Net income"), ("net_income_parent", "  of which: to parent"),
            ("eps_basic", "EPS (₱)")]
    rows = [("", ["Q3'24", "Q3'23", "9M'24", "9M'23"])]
    for acct, label in want:
        li = inc.get(acct)
        if not li:
            continue
        vals = []
        for pkey in config.INCOME_PERIODS:
            x = li["values"].get(pkey)
            vals.append(f"{x:.2f}" if acct == "eps_basic" and x is not None else (money(x / 1000) if x is not None else "—"))
        rows.append((label, vals))
    _grid(doc, rows, first_w=2.3)

    para(doc, "Valuation summary (DCF)", size=10.5, bold=True, color=NAVY, before=8, after=3)
    vr = [("", ["₱m / x"])]
    for label, val, fmt in [
        ("Enterprise value", v["enterprise_value"] / 1000, "m"),
        ("Less: net debt", -v["net_debt"] / 1000, "m"),
        ("Less: NCI", -v["nci"] / 1000, "m"),
        ("Less: preferred", -v["preferred"] / 1000, "m"),
        ("Equity value", v["common_equity_value"] / 1000, "m"),
        ("Fair value / share (₱)", v["fair_value_per_share"], "p"),
        ("WACC", mv["wacc"]["wacc"], "pct"),
        ("Terminal growth", mv["drivers"]["terminal_growth"], "pct"),
        ("Beta", ctx["market"].get("beta"), "x"),
    ]:
        if fmt == "m":
            s = money(val)
        elif fmt == "p":
            s = f"₱{val:.2f}"
        elif fmt == "pct":
            s = f"{val:.1%}"
        else:
            s = f"{val}"
        vr.append((label, [s]))
    _grid(doc, vr, first_w=2.6)

    doc.add_picture(str(charts.valuation_bar(mv)), width=Inches(3.2))


def _narrative(doc, ctx, nar):
    doc.add_page_break()
    heading(doc, "Earnings recap")
    para(doc, nar["earnings_recap"], size=10)
    heading(doc, "Sector context")
    para(doc, nar["sector_context"], size=10)
    heading(doc, "Outlook")
    para(doc, nar["outlook"], size=10)
    if ctx.get("news", {}).get("bullets"):
        heading(doc, f"Recent developments (live, as of {ctx['news'].get('as_of','')})")
        for b in ctx["news"]["bullets"]:
            bp = doc.add_paragraph(style="List Bullet"); run(bp, b, size=9.5)
    heading(doc, "Key risks")
    for r in nar["risks"]:
        bp = doc.add_paragraph(style="List Bullet")
        lead, _, rest = r.partition("—")
        run(bp, lead.strip() + (" — " if rest else ""), size=9.5, bold=True)
        run(bp, rest.strip(), size=9.5)


def _financial_highlights(doc, ctx):
    doc.add_page_break()
    heading(doc, "Financial highlights")
    p = ctx["model_values"]["projections"]
    yrs = [2024, 2025, 2026, 2027]
    hdr = [f"{y}{'E' if y<=2027 else 'F'}" for y in yrs]
    rows = [("₱m", hdr)]
    for key, label in [("revenue", "Revenue"), ("ebit", "EBIT"), ("nopat", "NOPAT"),
                       ("da", "D&A"), ("capex", "CapEx"), ("fcf", "Free cash flow"),
                       ("net_income", "Net income (parent)")]:
        rows.append((label, [money(p[key][str(y)] / 1000) for y in yrs]))
    rows.append(("EPS (₱)", [f"{p['eps'][str(y)]:.2f}" for y in yrs]))
    _grid(doc, rows, first_w=2.3)
    para(doc, "2024E is FY2024 estimated from 9M-2024 actuals (annualized); 2025E–2027F are model "
              "forecasts off configurable assumptions. Figures are illustrative.", size=8, color=GREY, before=4)
    doc.add_picture(str(charts.revenue_fcf(ctx["model_values"])), width=Inches(4.2))


def _disclaimer(doc, ctx):
    doc.add_page_break()
    heading(doc, "Rating definitions & disclaimer")
    para(doc, "Ratings (12-month total return): BUY > 10% · HOLD 0–10% · UNDERPERFORM < 0%.",
         size=9, bold=True)
    para(doc,
         "This document is an automated draft produced by a demonstration research-automation "
         "pipeline. Financial figures are extracted from the company's PSE filing and validated "
         "for internal consistency; market data is as of the stated date; forecasts derive from "
         "configurable assumptions and are illustrative, not house views. The terminal value uses "
         "a standard perpetuity, an approximation for a partly finite concession. This is not "
         "investment advice, an offer, or a solicitation. Sources: company SEC Form 17-Q (Q3 2024); "
         "public market data.", size=8.5, color=GREY)


# ── assemble ──────────────────────────────────────────────────────────────
def build() -> str:
    ctx = load_context()
    nar = narrative.generate(ctx)

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.6)

    _front(doc, ctx, nar)
    _key_highlights(doc, ctx)
    _narrative(doc, ctx, nar)
    _financial_highlights(doc, ctx)
    _disclaimer(doc, ctx)

    out = config.output_dir_for() / "mwc_research_note.docx"
    doc.save(out)
    return str(out)


if __name__ == "__main__":
    from pipeline.logger import console
    path = build()
    console.print(f"[bold]Research note[/] → {path}")
